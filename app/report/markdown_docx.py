"""Render determinista de Markdown (GitHub) a Word (.docx).

Soporta encabezados, párrafos con **negrita**/*cursiva*/`código`, listas, citas
(>), tablas y **imágenes** `![alt](ruta)` (para embeber gráficos). Usado por la
suite de informes para convertir la salida en Markdown del agente a Word.
"""
from __future__ import annotations
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AZUL = "0B3C5D"
_INLINE = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`")
_IMG = re.compile(r"^!\[(.*?)\]\((.+?)\)\s*$")


def _runs(p, text):
    pos = 0
    for m in _INLINE.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        if m.group(1) is not None:
            p.add_run(m.group(1)).bold = True
        elif m.group(2) is not None:
            p.add_run(m.group(2)).italic = True
        elif m.group(3) is not None:
            r = p.add_run(m.group(3)); r.font.name = "Consolas"
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


def _shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    s = OxmlElement("w:shd"); s.set(qn("w:val"), "clear"); s.set(qn("w:fill"), hexc)
    tcPr.append(s)


def _table(doc, block):
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in block]
    header, body = cells[0], cells[2:]
    n = len(header)
    t = doc.add_table(rows=1, cols=n); t.style = "Table Grid"; t.autofit = True
    for j, h in enumerate(header):
        c = t.rows[0].cells[j]; c.paragraphs[0].text = ""; _runs(c.paragraphs[0], h)
        for run in c.paragraphs[0].runs:
            run.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(c, AZUL)
    for row in body:
        row = (row + [""] * n)[:n]
        rc = t.add_row().cells
        for j, v in enumerate(row):
            rc[j].paragraphs[0].text = ""; _runs(rc[j].paragraphs[0], v)


def render_markdown(md: str, out_path: str, img_width_in: float = 5.3) -> str:
    """Convierte `md` a un .docx en `out_path`. Devuelve la ruta."""
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    lines = md.split("\n")
    i = 0
    while i < len(lines):
        ln = lines[i].rstrip()
        if not ln.strip():
            i += 1; continue
        mi = _IMG.match(ln.strip())
        if mi:
            path = mi.group(2)
            if os.path.exists(path):
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(path, width=Inches(img_width_in))
                cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = cap.add_run(mi.group(1)); r.italic = True
                r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i += 1; continue
        if ln.lstrip().startswith("|"):
            blk = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                blk.append(lines[i]); i += 1
            if len(blk) >= 2:
                _table(doc, blk)
            continue
        if set(ln.strip()) <= set("-") and len(ln.strip()) >= 3:
            i += 1; continue
        if ln.startswith("#### "):
            _runs(doc.add_paragraph(style="Heading 3"), ln[5:])
        elif ln.startswith("### "):
            _runs(doc.add_paragraph(style="Heading 2"), ln[4:])
        elif ln.startswith("## "):
            _runs(doc.add_paragraph(style="Heading 1"), ln[3:])
        elif ln.startswith("# "):
            _runs(doc.add_paragraph(style="Title"), ln[2:])
        elif ln.lstrip().startswith(">"):
            txt = ln.lstrip()[1:].strip()
            if txt:
                p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.25)
                _runs(p, txt)
                for r in p.runs:
                    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        elif re.match(r"^\s*[-*] ", ln):
            _runs(doc.add_paragraph(style="List Bullet"), re.sub(r"^\s*[-*] ", "", ln))
        elif re.match(r"^\s*\d+\. ", ln):
            _runs(doc.add_paragraph(style="List Number"), re.sub(r"^\s*\d+\. ", "", ln))
        else:
            _runs(doc.add_paragraph(), ln)
        i += 1
    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc.save(out_path)
    return out_path
