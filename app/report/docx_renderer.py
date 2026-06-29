"""Renderer determinista: Informe (dict validado) -> .docx profesional.

Sin IA aquí. Paleta institucional sobria, portada, secciones con tablas
sombreadas, riesgos coloreados por nivel, fuentes como hipervínculos y pie con
numeración de páginas. Robusto a campos faltantes (omite secciones vacías).
"""
from __future__ import annotations

import datetime as dt
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.config import Settings

# Paleta institucional.
AZUL = "0B3C5D"
ACENTO = "1D6FA3"
GRIS = "444444"
NIVEL_COLOR = {"alto": "C0392B", "medio": "E67E22", "bajo": "27AE60"}
PRIORIDAD_COLOR = {"P1": "C0392B", "P2": "E67E22", "P3": "27AE60"}
BLANCO = "FFFFFF"


# --------------------------------------------------------------------------- #
# Helpers OOXML
# --------------------------------------------------------------------------- #
def _set_cell_shading(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _cell_text(cell, text: str, *, bold=False, color=GRIS, size=10.5, align=None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(str(text) if text is not None else "")
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def _add_page_number_footer(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = ""
    run = p.add_run("Página ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(GRIS)
    _add_field(p, "PAGE")
    run2 = p.add_run(" de ")
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor.from_string(GRIS)
    _add_field(p, "NUMPAGES")
    nota = p.add_run("    ·    CONFIDENCIAL — Uso oficial restringido")
    nota.font.size = Pt(8)
    nota.font.color.rgb = RGBColor.from_string(GRIS)


def _add_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    fldBegin = OxmlElement("w:fldChar")
    fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_code} "
    fldEnd = OxmlElement("w:fldChar")
    fldEnd.set(qn("w:fldCharType"), "end")
    run._r.append(fldBegin)
    run._r.append(instr)
    run._r.append(fldEnd)


def _heading(doc, text, level=1, color=AZUL):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.color.rgb = RGBColor.from_string(color)
    return h


def add_hyperlink(paragraph, url: str, text: str, color=ACENTO):
    """Inserta un hipervínculo clicable en un párrafo."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text or url
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _body(doc, text, color=GRIS, size=10.5, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.italic = italic
    return p


# --------------------------------------------------------------------------- #
# Bloques de contenido
# --------------------------------------------------------------------------- #
def _kpi_table(doc, kpis):
    if not kpis:
        return
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for c, txt in zip(hdr, ("Indicador", "Valor", "Nota")):
        _set_cell_shading(c, AZUL)
        _cell_text(c, txt, bold=True, color=BLANCO)
    for k in kpis:
        row = table.add_row().cells
        _cell_text(row[0], k.get("label", ""), bold=True, color=AZUL)
        _cell_text(row[1], k.get("valor", ""))
        _cell_text(row[2], k.get("nota") or "")
    doc.add_paragraph()


def _data_table(doc, tabla):
    cols = tabla.get("columnas") or []
    filas = tabla.get("filas") or []
    if not cols:
        return
    if tabla.get("titulo"):
        _body(doc, tabla["titulo"], color=ACENTO, size=11).runs[0].bold = True
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    for c, txt in zip(table.rows[0].cells, cols):
        _set_cell_shading(c, AZUL)
        _cell_text(c, txt, bold=True, color=BLANCO)
    for fila in filas:
        row = table.add_row().cells
        for i, _ in enumerate(cols):
            val = fila[i] if i < len(fila) else ""
            _cell_text(row[i], val)
    doc.add_paragraph()


def _risk_block(doc, riesgos):
    for r in riesgos:
        nivel = (r.get("nivel") or "medio").lower()
        color = NIVEL_COLOR.get(nivel, ACENTO)
        p = doc.add_paragraph()
        tag = p.add_run(f"  {nivel.upper()}  ")
        tag.bold = True
        tag.font.size = Pt(9)
        tag.font.color.rgb = RGBColor.from_string(BLANCO)
        _shade_run(tag, color)
        titulo = p.add_run(f"  {r.get('riesgo', '')}")
        titulo.bold = True
        titulo.font.size = Pt(11)
        titulo.font.color.rgb = RGBColor.from_string(AZUL)
        if r.get("descripcion"):
            _body(doc, r["descripcion"])
        if r.get("evidencia_datos"):
            _body(doc, f"Evidencia: {r['evidencia_datos']}", color=ACENTO, size=9.5, italic=True)


def _shade_run(run, hex_color: str) -> None:
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    rPr.append(shd)


def mapa_puntos_ciegos(zonas: list[dict], out_png: str) -> str | None:
    """Gráfico de burbujas de puntos ciegos: X=MMI, Y=habitantes (log), tamaño=
    déficit de reportes, color=crítica(rojo)/silenciosa(naranja)/visible(verde).
    Devuelve la ruta del PNG, o None si no hay matplotlib o datos."""
    if not zonas:
        return None
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except Exception:  # noqa: BLE001 — sin matplotlib, se omite el gráfico
        return None

    xs, ys, sizes, colors = [], [], [], []
    etiquetas = []
    for z in zonas:
        mmi = float(z.get("mmi") or 0)
        hab = max(1, int(z.get("habitantes") or 1))
        deficit = max(0.0, float(z.get("reportes_esperados") or 0) - int(z.get("reportes_observados") or 0))
        critica = bool(z.get("critica"))
        cobertura = float(z.get("cobertura") or 0)
        xs.append(mmi); ys.append(hab)
        sizes.append(40 + min(900, deficit * 4))
        colors.append("#C0392B" if critica else ("#E67E22" if cobertura < 0.6 else "#27AE60"))
        if critica:
            etiquetas.append((mmi, hab, z.get("localidad", "")))

    fig, ax = plt.subplots(figsize=(7.2, 4.6))
    ax.scatter(xs, ys, s=sizes, c=colors, alpha=0.7, edgecolors="white", linewidths=0.6)
    ax.set_yscale("log")
    ax.set_xlabel("Intensidad sísmica (MMI)")
    ax.set_ylabel("Habitantes (escala log)")
    ax.set_title("Mapa de puntos ciegos — intensidad × población × déficit de reportes",
                 fontsize=10, color="#0B3C5D")
    ax.grid(True, which="both", linestyle=":", alpha=0.3)
    for mmi, hab, nombre in etiquetas[:12]:
        ax.annotate(nombre, (mmi, hab), fontsize=7, color="#444444",
                    xytext=(4, 4), textcoords="offset points")
    # Leyenda manual.
    from matplotlib.lines import Line2D
    leg = [Line2D([0], [0], marker="o", color="w", label=lbl, markerfacecolor=col, markersize=8)
           for lbl, col in (("Crítica", "#C0392B"), ("Silenciosa", "#E67E22"), ("Visible", "#27AE60"))]
    ax.legend(handles=leg, loc="lower right", fontsize=8, framealpha=0.9)
    fig.tight_layout()
    fig.savefig(out_png, dpi=130)
    plt.close(fig)
    return out_png


def _seccion_puntos_ciegos(doc, informe):
    zonas = informe.get("zonas_ciegas") or []
    resumen = informe.get("resumen_puntos_ciegos") or ""
    if not zonas and not resumen:
        return
    _heading(doc, "Zonas Silenciosas y Puntos Ciegos (cruce intensidad × reportes)", 2)
    # Advertencia metodológica.
    adv = doc.add_paragraph()
    advr = adv.add_run(
        "Advertencia metodológica: la baja o nula recepción de reportes NO significa ausencia de "
        "daño. Estas zonas se tratan como ZONA DE ATENCIÓN (probable caída de conectividad/energía o "
        "cercanía al epicentro) y se priorizan para verificación en terreno."
    )
    advr.italic = True
    advr.font.size = Pt(10)
    advr.font.color.rgb = RGBColor.from_string("C0392B")
    if resumen:
        _body(doc, resumen)

    # Tabla de zonas (críticas resaltadas).
    criticas = [z for z in zonas if z.get("critica")]
    orden = sorted(zonas, key=lambda z: (not z.get("critica"), -float(z.get("indice_punto_ciego") or 0)))
    cols = ["Localidad", "Estado", "MMI", "Habitantes", "Reportes", "Esperados", "Cobertura", "Prioridad"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    for c, txt in zip(table.rows[0].cells, cols):
        _set_cell_shading(c, AZUL)
        _cell_text(c, txt, bold=True, color=BLANCO)
    for z in orden:
        row = table.add_row().cells
        vals = [
            z.get("localidad", ""), z.get("estado") or "—", z.get("mmi", ""),
            f"{int(z.get('habitantes') or 0):,}".replace(",", "."),
            z.get("reportes_observados", 0), round(float(z.get("reportes_esperados") or 0)),
            f"{float(z.get('cobertura') or 0):.0%}", z.get("prioridad") or ("P1" if z.get("critica") else "—"),
        ]
        for i, v in enumerate(vals):
            _cell_text(row[i], v)
        if z.get("critica"):
            _set_cell_shading(row[0], "F9D7D3")  # resalte suave de fila crítica
    doc.add_paragraph()

    # Mapa de burbujas.
    try:
        import os
        png = os.path.join(Settings().reports_out_dir, "_mapa_puntos_ciegos.png")
        if mapa_puntos_ciegos(orden, png):
            doc.add_picture(png, width=Inches(6.2))
            doc.add_paragraph()
    except Exception:  # noqa: BLE001
        pass

    # Recomendaciones de verificación para zonas críticas.
    if criticas:
        _body(doc, "Verificación prioritaria en terreno (zonas críticas):", color=AZUL, size=11).runs[0].bold = True
        for z in criticas[:10]:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(
                f"{z.get('localidad','')} ({z.get('estado') or 's/d'}) — MMI {z.get('mmi')}, "
                f"~{int(z.get('habitantes') or 0):,} hab, {z.get('reportes_observados',0)} reportes."
            )
            r.font.color.rgb = RGBColor.from_string(GRIS)
    # Nota de sin_match (zonas MMI sin geografía resuelta).
    sin_match = informe.get("puntos_ciegos_sin_match") or []
    if sin_match:
        _body(doc, "Sin geolocalizar (revisar manualmente): " + ", ".join(sin_match[:20]),
              color=GRIS, size=9, italic=True)


# --------------------------------------------------------------------------- #
# Secciones (en funciones para poder ordenarlas como informe EJECUTIVO)
# --------------------------------------------------------------------------- #
def _sec_hallazgos(doc, informe):
    hall = informe.get("hallazgos_clave") or []
    if not hall:
        return
    _heading(doc, "Hallazgos clave", 1)
    for h in hall:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(str(h))
        run.font.size = Pt(11.5)
        run.font.color.rgb = RGBColor.from_string(GRIS)
    doc.add_paragraph()


def _sec_recomendaciones(doc, informe):
    recs = informe.get("recomendaciones") or []
    if not recs:
        return
    _heading(doc, "Recomendaciones priorizadas", 1)
    orden = {"P1": 0, "P2": 1, "P3": 2}
    recs = sorted(recs, key=lambda r: orden.get(r.get("prioridad", "P3"), 9))
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for c, txt in zip(table.rows[0].cells, ("Prioridad", "Acción", "Responsable", "Plazo")):
        _set_cell_shading(c, AZUL)
        _cell_text(c, txt, bold=True, color=BLANCO)
    for r in recs:
        row = table.add_row().cells
        pr = r.get("prioridad", "P3")
        _cell_text(row[0], pr, bold=True, color=BLANCO, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_shading(row[0], PRIORIDAD_COLOR.get(pr, ACENTO))
        _cell_text(row[1], r.get("accion", ""))
        _cell_text(row[2], r.get("responsable_sugerido") or "")
        _cell_text(row[3], r.get("plazo") or "")
    doc.add_paragraph()


def _sec_conclusiones(doc, informe):
    concl = informe.get("conclusiones") or ""
    if not concl:
        return
    _heading(doc, "Conclusiones", 1)
    _body(doc, concl, size=11)


def _sec_plan_comunicaciones(doc, informe):
    plan = informe.get("plan_comunicaciones") or {}
    mensajes = plan.get("mensajes_alerta") or []
    lineamientos = plan.get("lineamientos") or []
    if not (plan.get("objetivo") or plan.get("voceria") or mensajes or lineamientos):
        return
    _heading(doc, "Plan de comunicaciones", 1)
    if plan.get("objetivo"):
        _body(doc, plan["objetivo"], size=11)
    if plan.get("voceria"):
        p = _body(doc, f"Vocería oficial: {plan['voceria']}", color=AZUL)
        p.runs[0].bold = True
    if mensajes:
        _body(doc, "Mensajes de alerta a difundir:", color=ACENTO, size=11).runs[0].bold = True
        cols = ["Canal", "Audiencia", "Mensaje", "Momento"]
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = "Table Grid"
        for c, txt in zip(table.rows[0].cells, cols):
            _set_cell_shading(c, AZUL)
            _cell_text(c, txt, bold=True, color=BLANCO)
        for m in mensajes:
            row = table.add_row().cells
            _cell_text(row[0], m.get("canal", ""), bold=True, color=AZUL)
            _cell_text(row[1], m.get("audiencia", ""))
            _cell_text(row[2], m.get("mensaje", ""))
            _cell_text(row[3], m.get("momento") or "")
        doc.add_paragraph()
    if lineamientos:
        _body(doc, "Lineamientos de comunicación:", color=ACENTO, size=11).runs[0].bold = True
        for li in lineamientos:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(str(li))
            r.font.color.rgb = RGBColor.from_string(GRIS)
    doc.add_paragraph()


def _sec_resumen(doc, informe):
    if informe.get("resumen_ejecutivo"):
        _heading(doc, "Resumen ejecutivo", 2)
        _body(doc, informe["resumen_ejecutivo"], size=11)


def _sec_panorama(doc, informe):
    panorama = informe.get("panorama_datos") or {}
    if panorama.get("narrativa") or panorama.get("kpis") or panorama.get("tablas"):
        _heading(doc, "Panorama de datos", 2)
        if panorama.get("narrativa"):
            _body(doc, panorama["narrativa"])
        _kpi_table(doc, panorama.get("kpis") or [])
        for tabla in panorama.get("tablas") or []:
            _data_table(doc, tabla)


def _sec_riesgos(doc, informe):
    if informe.get("analisis_riesgos"):
        _heading(doc, "Análisis de riesgos", 2)
        _risk_block(doc, informe["analisis_riesgos"])


def _sec_predicciones(doc, informe):
    if informe.get("predicciones"):
        _heading(doc, "Predicciones", 2)
        for p in informe["predicciones"]:
            _body(doc, f"[{p.get('horizonte', '')}] {p.get('prediccion', '')}", color=AZUL).runs[0].bold = True
            if p.get("supuestos"):
                _body(doc, f"Supuestos: {p['supuestos']}", size=9.5, italic=True)


def _sec_casos(doc, informe):
    if informe.get("casos_analogos"):
        _heading(doc, "Casos análogos internacionales", 2)
        for c in informe["casos_analogos"]:
            _body(doc, f"{c.get('pais', '')} — {c.get('evento', '')}", color=ACENTO, size=11).runs[0].bold = True
            if c.get("leccion"):
                _body(doc, f"Lección: {c['leccion']}")
            if c.get("aplicacion_venezuela"):
                _body(doc, f"Aplicación a Venezuela: {c['aplicacion_venezuela']}")
            if c.get("fuente"):
                fp = doc.add_paragraph()
                add_hyperlink(fp, c["fuente"], f"Fuente: {c['fuente']}")


def _sec_protocolos(doc, informe):
    if informe.get("protocolos_recomendados"):
        _heading(doc, "Protocolos recomendados", 2)
        for p in informe["protocolos_recomendados"]:
            _body(doc, f"{p.get('organismo', '')} — {p.get('protocolo', '')}", color=ACENTO, size=11).runs[0].bold = True
            if p.get("resumen"):
                _body(doc, p["resumen"])
            if p.get("fuente"):
                fp = doc.add_paragraph()
                add_hyperlink(fp, p["fuente"], f"Fuente: {p['fuente']}")


def _sec_acciones(doc, informe):
    if informe.get("acciones_preventivas"):
        _heading(doc, "Acciones preventivas", 2)
        for a in informe["acciones_preventivas"]:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(a.get("accion", ""))
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(AZUL)
            if a.get("justificacion"):
                _body(doc, f"  Justificación: {a['justificacion']}", size=9.5, italic=True)


def _sec_indicadores(doc, informe):
    if informe.get("indicadores_nuevos"):
        _heading(doc, "Indicadores nuevos propuestos", 2)
        _data_table(doc, {
            "titulo": "",
            "columnas": ["Indicador", "Definición", "Cómo calcularlo"],
            "filas": [[i.get("indicador", ""), i.get("definicion", ""), i.get("como_calcularlo", "")]
                      for i in informe["indicadores_nuevos"]],
        })


def _sec_fuentes(doc, informe):
    if informe.get("fuentes"):
        _heading(doc, "Fuentes", 2)
        for n, f in enumerate(informe["fuentes"], 1):
            p = doc.add_paragraph()
            p.add_run(f"{n}. ").bold = True
            if f.get("url"):
                add_hyperlink(p, f["url"], f.get("titulo") or f["url"])
            else:
                p.add_run(f.get("titulo") or "")


# --------------------------------------------------------------------------- #
# Render principal
# --------------------------------------------------------------------------- #
def render_informe(informe: dict, out_path: str | None = None) -> str:
    informe = informe or {}
    meta = informe.get("meta", {}) or {}

    if out_path is None:
        s = Settings()
        tipo = (meta.get("tipo_reporte") or "informe").replace(" ", "_")
        stamp = dt.datetime.now().strftime("%Y%m%d_%H%M")
        out_path = os.path.join(s.reports_out_dir, f"informe_{tipo}_{stamp}.docx")
    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)

    doc = Document()
    # Estilo base del documento.
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(GRIS)

    _add_page_number_footer(doc.sections[0])

    # 2) PORTADA
    for _ in range(4):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run(meta.get("titulo") or "Informe Estratégico — Terremoto 24J")
    tr.bold = True
    tr.font.size = Pt(24)
    tr.font.color.rgb = RGBColor.from_string(AZUL)

    cl = doc.add_paragraph()
    cl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    clr = cl.add_run(meta.get("clasificacion") or "CONFIDENCIAL — Uso oficial restringido")
    clr.bold = True
    clr.font.size = Pt(11)
    clr.font.color.rgb = RGBColor.from_string("C0392B")

    doc.add_paragraph()
    for label, val in (
        ("Evento", "Terremoto 24J — Venezuela"),
        ("Tipo de informe", meta.get("tipo_reporte", "")),
        ("Ventana de datos", meta.get("ventana_datos", "")),
        ("Fecha de generación", meta.get("fecha_generacion", "")),
        ("Preparado por", meta.get("preparado_por", "Copiloto Estratégico IA — VenApp / Línea 58")),
    ):
        if not val:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = p.add_run(f"{label}: ")
        rl.bold = True
        rl.font.color.rgb = RGBColor.from_string(AZUL)
        rv = p.add_run(str(val))
        rv.font.color.rgb = RGBColor.from_string(GRIS)
    doc.add_page_break()

    # ===================== FRENTE EJECUTIVO (primeras páginas) ============== #
    # Para la Presidenta: lo más importante primero, lectura rápida.
    # 1) Hallazgos clave  ->  2) Recomendaciones  ->  3) Conclusiones
    _sec_hallazgos(doc, informe)
    _sec_recomendaciones(doc, informe)
    _sec_conclusiones(doc, informe)
    _sec_plan_comunicaciones(doc, informe)

    # ===================== DESARROLLO (detalle, después) =================== #
    doc.add_page_break()
    _heading(doc, "Desarrollo", 1)
    _sec_resumen(doc, informe)
    _sec_panorama(doc, informe)
    _seccion_puntos_ciegos(doc, informe)
    _sec_riesgos(doc, informe)
    _sec_predicciones(doc, informe)
    _sec_casos(doc, informe)
    _sec_protocolos(doc, informe)
    _sec_acciones(doc, informe)
    _sec_indicadores(doc, informe)
    _sec_fuentes(doc, informe)

    # NOTA PII
    if informe.get("notas_pii"):
        doc.add_paragraph()
        _body(doc, informe["notas_pii"], color=GRIS, size=9, italic=True)

    doc.save(out_path)
    return out_path
